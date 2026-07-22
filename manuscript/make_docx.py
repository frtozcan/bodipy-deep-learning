"""make_docx.py - manuscript/draft.md -> manuscript/BODIPY_manuscript.docx
Markdown başlık/paragraf/tablo/liste yapısını Word'e çevirir; results/ figürlerini gömer.
Çalıştır: python manuscript/make_docx.py
"""
import os, re, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = 'Times New Roman'
BODY_PT, TABLE_PT, CAPTION_PT = 12, 10, 10

def _force_font(style, name=FONT):
    """Stilin ascii/hAnsi/eastAsia fontlarini birlikte ayarla (Word aksi halde uygulamaz)."""
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(attr), name)

def setup_styles(doc):
    st = doc.styles['Normal']
    _force_font(st); st.font.size = Pt(BODY_PT); st.font.color.rgb = RGBColor(0, 0, 0)
    for nm in ['Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5']:
        try:
            h = doc.styles[nm]
        except KeyError:
            continue
        _force_font(h)
        h.font.color.rgb = RGBColor(0, 0, 0)      # basliklar SIYAH
        h.font.bold = True
        if nm == 'Title':
            h.font.size = Pt(16)
        elif nm == 'Heading 1':
            h.font.size = Pt(14)
        else:
            h.font.size = Pt(12)

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
# Kullanim: python make_docx.py [draft.md] [BODIPY_manuscript.docx] [dil: en|tr]
_md = sys.argv[1] if len(sys.argv) > 1 else 'draft.md'
_out = sys.argv[2] if len(sys.argv) > 2 else 'BODIPY_manuscript.docx'
LANG = sys.argv[3] if len(sys.argv) > 3 else 'en'
MD = os.path.join(HERE, _md)
OUT = os.path.join(HERE, _out)
RESULTS = os.path.join(ROOT, 'results')

# Figür yerleşimi: bölüm başlığı sonrası eklenecek görseller
FIGS = {
    'Fig1': ('multiseed_R2.png', 'Figure 1. Architecture comparison on the internal test set: R² per property (mean ± s.d., 5 seeds).'),
    'Fig2': ('extpool_parity.png', 'Figure 2. Pooled external validation: predicted vs measured properties for 190 dyes from twelve held-out publications.'),
    'Fig3a': ('probe_QY.png', 'Figure 3a. In-silico meso-substituent probe: predicted ΦF (5-seed ensemble).'),
    'Fig3b': ('probe_heavy_atom.png', 'Figure 3b. Predicted heavy-atom series (4-halophenyl at meso).'),
    'Fig4': ('attention_maps.png', 'Figure 4. GATv2 attention-derived atom importance (5-seed ensemble).'),
    'Fig5': ('shap_summary.png', 'Figure 5. SHAP feature attribution for ΦF (MLP, 5-seed ensemble).'),
    'Fig6': ('qy_bimodal.png', 'Figure 6. Quantum-yield parametrisations compared at both evaluation levels (MLP, 5 seeds).'),
    'Fig7': ('solvent_ablation.png', 'Figure 7. Solvent representation ablation (MLP, multi-task, 5 seeds).'),
    'FigS1': ('eda_bodipy_dist.png', 'Figure S1. Target distributions of the BODIPY subset.'),
    'FigS2': ('lopo_box.png', 'Figure S2. Error distribution across individual held-out publications (leave-one-publication-out).'),
    'FigS3': ('external_parity.png', 'Figure S3. Parity plots on the eight-dye independent synthetic study.'),
}

FIGS_TR = {
    'Fig1': ('multiseed_R2.png', 'Şekil 1. Mimarilerin iç test kümesindeki karşılaştırması: özellik başına R² (ortalama ± s.s., 5 tohum).'),
    'Fig2': ('external_parity.png', 'Şekil 2. Harici doğrulama: sekiz bağımsız BODIPY boyası için tahmin edilen ve ölçülen özellikler.'),
    'Fig3a': ('probe_QY.png', 'Şekil 3a. In-siliko mezo-sübstitüent taraması: tahmin edilen ΦF (5 tohumlu topluluk).'),
    'Fig3b': ('probe_heavy_atom.png', 'Şekil 3b. Tahmin edilen ağır atom serisi (mezo konumunda 4-halofenil).'),
    'Fig4': ('attention_maps.png', 'Şekil 4. GATv2 dikkat temelli atom önem haritaları (5 tohumlu topluluk).'),
    'Fig5': ('shap_summary.png', 'Şekil 5. ΦF için SHAP öznitelik katkıları (MLP, 5 tohumlu topluluk).'),
    'FigS1': ('eda_bodipy_dist.png', 'Şekil E1. BODIPY alt kümesinin hedef dağılımları.'),
}
if LANG == 'tr':
    FIGS = FIGS_TR

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')

def add_runs(par, text):
    """**bold**, *italic*, `code` işaretlemesini Word run'larına çevir."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            par.add_run(part[2:-2]).bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = par.add_run(part[1:-1]); r.font.name = FONT; r.font.size = Pt(BODY_PT)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            par.add_run(part[1:-1]).italic = True
        else:
            par.add_run(part)

def add_figure(doc, key):
    fname, caption = FIGS[key]
    path = os.path.join(RESULTS, fname)
    if not os.path.exists(path):
        print('  [atlandi, yok]', fname); return
    doc.add_picture(path, width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(CAPTION_PT)
    r.font.name = FONT
    print('  [figur]', fname)

def add_table(doc, rows):
    header = [c.strip() for c in rows[0].strip('|').split('|')]
    body = [[c.strip() for c in r.strip('|').split('|')] for r in rows[2:]]
    t = doc.add_table(rows=1, cols=len(header)); t.style = 'Table Grid'   # renksiz, duz cerceve
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]; cell.text = ''
        add_runs(cell.paragraphs[0], h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True; run.font.size = Pt(TABLE_PT)
                run.font.name = FONT; run.font.color.rgb = RGBColor(0, 0, 0)
    for row in body:
        cells = t.add_row().cells
        for i, c in enumerate(row[:len(header)]):
            cells[i].text = ''
            add_runs(cells[i].paragraphs[0], c)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(TABLE_PT)
                    run.font.name = FONT; run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

def main():
    lines = open(MD, encoding='utf-8').read().split('\n')
    doc = Document()
    setup_styles(doc)

    i, n = 0, len(lines)
    skip_block = False       # taslak notlarını (blockquote) atla
    while i < n:
        ln = lines[i]; s = ln.strip()

        # Markdown tablosu
        if s.startswith('|') and i + 1 < n and set(lines[i+1].strip()) <= set('|-: '):
            blk = []
            while i < n and lines[i].strip().startswith('|'):
                blk.append(lines[i]); i += 1
            add_table(doc, blk); continue

        if s.startswith('> '):          # draft status notu -> atla
            i += 1; continue
        if s == '---':
            i += 1; continue

        if s.startswith('#'):
            lvl = len(s) - len(s.lstrip('#'))
            txt = s.lstrip('#').strip()
            if lvl == 1:
                h = doc.add_heading('', level=0); add_runs(h, txt)
            else:
                h = doc.add_heading('', level=min(lvl - 1, 4)); add_runs(h, txt)
            # bölüm sonrası figürler
            key = {'3.1': 'Fig1', '3.2': 'Fig2', '3.4': 'Fig3a',
                   '3.5': 'Fig4'}.get(txt.split()[0] if txt else '', None)
            i += 1
            continue

        if s.startswith('- ') or s.startswith('* '):
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, s[2:])
            i += 1; continue
        if re.match(r'^\d+\.\s', s):
            p = doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\d+\.\s', '', s))
            i += 1; continue

        if not s:
            i += 1; continue

        # normal paragraf: ardışık satırları birleştir
        buf = []
        while i < n and lines[i].strip() and not lines[i].strip().startswith(('#', '|', '- ', '> ')) \
                and lines[i].strip() != '---' and not re.match(r'^\d+\.\s', lines[i].strip()):
            buf.append(lines[i].strip()); i += 1
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, ' '.join(buf))
    return doc

if __name__ == '__main__':
    doc = main()
    # Figürleri sona, "Figures" bölümü olarak ekle (bölüm içine gömmek yerine
    # dergi formatına uygun şekilde toplu halde)
    doc.add_page_break()
    h = doc.add_heading('', level=1)
    add_runs(h, 'Şekiller' if LANG == 'tr' else 'Figures')
    for key in ['Fig1', 'Fig2', 'Fig3a', 'Fig3b', 'Fig4', 'Fig5', 'Fig6', 'Fig7',
                'FigS1', 'FigS2', 'FigS3']:
        if key not in FIGS:
            continue
        add_figure(doc, key)
        doc.add_paragraph()
    doc.save(OUT)
    print('\nKaydedildi:', OUT)
